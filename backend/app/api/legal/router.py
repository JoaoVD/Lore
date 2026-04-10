"""
app/api/legal/router.py
------------------------
Endpoints do Lore Jurídico — templates jurídicos pré-configurados.

Rotas:
  GET  /api/legal/templates
  GET  /api/legal/templates/{template_id}
  POST /api/projects/{project_id}/legal/generate
  GET  /api/projects/{project_id}/legal/documents
  POST /api/projects/{project_id}/legal/templates
  GET  /api/projects/{project_id}/legal/documents/{document_id}/export
  POST /api/projects/{project_id}/legal/documents/{document_id}/feedback
  GET  /api/projects/{project_id}/legal/report
  GET  /api/legal/tribunais
  POST /api/projects/{project_id}/legal/cases/{case_id}/sync-pje
"""

import io
import json
import os
import re
from collections import Counter, defaultdict
from datetime import datetime, timedelta
from typing import Optional

import anthropic as _anthropic
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from supabase import Client

from app.core.features import Features
from app.db.supabase import get_supabase
from auth.middleware import AuthUser, get_current_user

router = APIRouter()


class GenerateDocumentRequest(BaseModel):
    template_id: str
    variables: dict        # {"nome_cliente": "João Silva", ...}
    use_ai: bool = False   # se True, usa IA para preencher campos vazios


# ── Templates ─────────────────────────────────────────────────────────────────

_AREA_META = {
    "trabalhista":    {"name": "Trabalhista",     "icon": "⚖️"},
    "civel":          {"name": "Cível",            "icon": "📋"},
    "contratos":      {"name": "Contratos",        "icon": "📝"},
    "familia":        {"name": "Família",           "icon": "👨‍👩‍👧"},
    "previdenciario": {"name": "Previdenciário",   "icon": "🏛️"},
}


def _normalize_diamond(t: dict) -> dict:
    """Converte um diamond_template para o shape esperado pelo frontend."""
    area = t.get("area", "")
    meta = _AREA_META.get(area, {"name": area.title(), "icon": "📄"})
    return {
        **t,
        # frontend usa 'content' para preview; diamond usa 'structure' — expõe vazio
        "content": "",
        "description": f"Padrão Diamante — {meta['name']}",
        "is_default": True,
        "legal_template_categories": meta,
        "category_id": area,
    }


@router.get("/api/legal/templates")
async def list_templates(
    project_id: Optional[str] = None,
    category: Optional[str] = None,
    user: AuthUser = Depends(get_current_user),
    supabase: Client = Depends(get_supabase),
):
    """Lista todos os templates disponíveis (Padrão Diamante + legados + do projeto)."""
    if not Features.JURIDICO:
        raise HTTPException(status_code=404, detail="Feature não disponível")

    # --- Templates Padrão Diamante (globais + do projeto) ---
    diamond_query = supabase.table("diamond_templates").select("*")
    if project_id:
        diamond_query = diamond_query.or_(f"project_id.is.null,project_id.eq.{project_id}")
    else:
        diamond_query = diamond_query.is_("project_id", "null")
    diamond_result = diamond_query.order("created_at").execute()
    diamond_templates = [_normalize_diamond(t) for t in (diamond_result.data or [])]

    # --- Templates legados ---
    legacy_query = supabase.table("legal_templates") \
        .select("*, legal_template_categories(name, icon)")
    if project_id:
        legacy_query = legacy_query.or_(f"is_default.eq.true,project_id.eq.{project_id}")
    else:
        legacy_query = legacy_query.eq("is_default", True)
    if category:
        legacy_query = legacy_query.eq("legal_template_categories.name", category)
    legacy_result = legacy_query.order("created_at").execute()

    all_templates = diamond_templates + (legacy_result.data or [])
    return {"templates": all_templates}


@router.get("/api/legal/templates/{template_id}")
async def get_template(
    template_id: str,
    user: AuthUser = Depends(get_current_user),
    supabase: Client = Depends(get_supabase),
):
    """Retorna um template específico com todas as variáveis."""
    if not Features.JURIDICO:
        raise HTTPException(status_code=404, detail="Feature não disponível")

    result = supabase.table("legal_templates") \
        .select("*, legal_template_categories(name, icon)") \
        .eq("id", template_id) \
        .execute()

    if not result.data:
        raise HTTPException(status_code=404, detail="Template não encontrado")

    return result.data[0]


# ── Documentos ────────────────────────────────────────────────────────────────

@router.post("/api/projects/{project_id}/legal/generate")
async def generate_document(
    project_id: str,
    body: GenerateDocumentRequest,
    user: AuthUser = Depends(get_current_user),
    supabase: Client = Depends(get_supabase),
):
    """Gera um documento a partir de um template (Padrão Diamante ou legado)."""
    if not Features.JURIDICO:
        raise HTTPException(status_code=404, detail="Feature não disponível")

    from app.services.diamond_generator import generate_diamond_document

    # Tenta Padrão Diamante primeiro
    diamond = supabase.table("diamond_templates") \
        .select("*").eq("id", body.template_id).execute()

    if diamond.data:
        template_data = diamond.data[0]
        result = await generate_diamond_document(
            template=template_data,
            variables=body.variables,
            project_id=project_id,
            supabase=supabase,
        )
        content = result["content"]
    else:
        # Fallback: template legado com substituição simples + IA opcional
        legacy = supabase.table("legal_templates") \
            .select("*").eq("id", body.template_id).execute()
        if not legacy.data:
            raise HTTPException(status_code=404, detail="Template não encontrado")
        template_data = legacy.data[0]
        content = template_data["content"]
        for key, value in body.variables.items():
            content = content.replace(f"{{{{{key}}}}}", str(value))
        if body.use_ai:
            remaining = re.findall(r'\{\{(\w+)\}\}', content)
            if remaining:
                content = await _fill_with_ai(content, remaining, template_data["name"])

    first_var = next(iter(body.variables.values()), "Documento") if body.variables else "Documento"
    doc_name = f"{template_data['name']} — {first_var}"[:100]

    saved = supabase.table("legal_documents").insert({
        "project_id":  project_id,
        "template_id": body.template_id,
        "name":        doc_name,
        "content":     content,
        "variables":   body.variables,
        "created_by":  user.id,
    }).execute()

    return {
        "document_id": saved.data[0]["id"],
        "name":        doc_name,
        "content":     content,
        "variables_filled":     len(body.variables),
        "variables_remaining":  len(re.findall(r'\{\{(\w+)\}\}', content)),
    }


@router.post("/api/projects/{project_id}/legal/generate-stream")
async def generate_document_stream(
    project_id: str,
    body: GenerateDocumentRequest,
    user: AuthUser = Depends(get_current_user),
    supabase: Client = Depends(get_supabase),
):
    """
    Geração com streaming — o texto aparece em tempo real no frontend.
    Retorna Server-Sent Events (SSE).
    """
    if not Features.JURIDICO:
        raise HTTPException(status_code=404, detail="Feature não disponível")

    # Busca o template (Diamante ou legado)
    template_result = supabase.table("diamond_templates").select("*").eq("id", body.template_id).execute()
    if not template_result.data:
        template_result = supabase.table("legal_templates").select("*").eq("id", body.template_id).execute()
    if not template_result.data:
        raise HTTPException(status_code=404, detail="Template não encontrado")

    template_data = template_result.data[0]
    variables     = body.variables
    structure     = template_data.get("structure", {})
    tone_guide    = template_data.get("tone_guide", "")

    # Estilo aprendido do escritório
    style = None
    try:
        style_result = supabase.table("office_style") \
            .select("style_summary").eq("project_id", project_id).execute()
        if style_result.data:
            style = style_result.data[0].get("style_summary", "")
    except Exception:
        pass

    vars_text = "\n".join(
        f"- {k}: {v}" for k, v in variables.items() if v and str(v).strip()
    )
    style_block = f"\nESTILO DO ESCRITÓRIO:\n{style}" if style else ""

    sections_instructions = ""
    for sec_name, sec_data in structure.items():
        if isinstance(sec_data, dict) and sec_data.get("instrucao"):
            sections_instructions += f"\n## {sec_name.upper()}\n{sec_data['instrucao']}\n"

    prompt = (
        f"Você é um advogado sênior brasileiro com 20 anos de experiência, "
        f"especialista em {template_data.get('area', 'direito')}.\n\n"
        f"Gere o documento jurídico completo \"{template_data.get('name')}\" com base nos dados abaixo.\n\n"
        f"DADOS DO CASO:\n{vars_text}\n\n"
        f"INSTRUÇÕES POR SEÇÃO:\n{sections_instructions}\n\n"
        f"TOM E ESTILO:\n{tone_guide}\n"
        f"{style_block}\n\n"
        f"REGRAS ABSOLUTAS:\n"
        f"1. Escrita EXTREMAMENTE formal, técnica e rebuscada\n"
        f"2. Use: \"ora Reclamante\", \"consoante dispõe\", \"à luz do que preceitua\", "
        f"\"importa salientar\", \"cumpre destacar\"\n"
        f"3. Cite legislação específica (CLT, CC, CPC, CF) quando relevante\n"
        f"4. Use {{{{variavel}}}} para dados não fornecidos — NUNCA invente\n"
        f"5. Gere o documento COMPLETO do início ao fim\n\n"
        f"Gere o documento agora:"
    )

    async def event_stream():
        client = _anthropic.AsyncAnthropic(api_key=os.getenv("ANTHROPIC_API_KEY", ""))
        full_text = ""
        try:
            async with client.messages.stream(
                model=os.getenv("CLAUDE_MODEL", "claude-sonnet-4-5"),
                max_tokens=4000,
                messages=[{"role": "user", "content": prompt}],
            ) as stream:
                async for text_chunk in stream.text_stream:
                    full_text += text_chunk
                    data = json.dumps({"type": "chunk", "text": text_chunk})
                    yield f"data: {data}\n\n"

            # Salva no banco após streaming completo
            first_val  = list(variables.values())[0] if variables else "Documento"
            doc_name   = f"{template_data['name']} — {first_val}"[:100]
            saved = supabase.table("legal_documents").insert({
                "project_id":  project_id,
                "template_id": body.template_id,
                "name":        doc_name,
                "content":     full_text,
                "variables":   variables,
                "created_by":  str(user.id),
            }).execute()
            doc_id = saved.data[0]["id"] if saved.data else None

            done_data = json.dumps({"type": "done", "document_id": doc_id, "name": doc_name})
            yield f"data: {done_data}\n\n"

        except Exception as exc:
            error_data = json.dumps({"type": "error", "message": str(exc)})
            yield f"data: {error_data}\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control":     "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


@router.post("/api/projects/{project_id}/legal/documents/{document_id}/feedback")
async def save_document_feedback(
    project_id: str,
    document_id: str,
    body: dict,
    user: AuthUser = Depends(get_current_user),
    supabase: Client = Depends(get_supabase),
):
    """
    Salva a edição do advogado para o feedback loop.
    A cada 5 feedbacks, o Lore atualiza o estilo do escritório automaticamente.
    """
    if not Features.JURIDICO:
        raise HTTPException(status_code=404, detail="Feature não disponível")

    from app.services.diamond_generator import save_feedback_and_learn

    original = body.get("original", "")
    edited   = body.get("edited", "")
    section  = body.get("section", "geral")

    if not edited or original == edited:
        return {"message": "Sem alterações detectadas."}

    await save_feedback_and_learn(
        project_id=project_id,
        document_id=document_id,
        section=section,
        original=original,
        edited=edited,
        user_id=user.id,
        supabase=supabase,
    )

    # Atualiza o conteúdo do documento salvo
    supabase.table("legal_documents").update({"content": edited}) \
        .eq("id", document_id).eq("project_id", project_id).execute()

    return {"message": "Feedback salvo. O Lore está aprendendo com sua edição."}


@router.get("/api/projects/{project_id}/legal/documents")
async def list_documents(
    project_id: str,
    user: AuthUser = Depends(get_current_user),
    supabase: Client = Depends(get_supabase),
):
    """Lista documentos gerados do projeto."""
    if not Features.JURIDICO:
        raise HTTPException(status_code=404, detail="Feature não disponível")

    result = supabase.table("legal_documents") \
        .select("id, name, created_at, template_id, legal_templates(name)") \
        .eq("project_id", project_id) \
        .order("created_at", desc=True) \
        .execute()

    return {"documents": result.data}


@router.post("/api/projects/{project_id}/legal/templates")
async def create_custom_template(
    project_id: str,
    body: dict,
    user: AuthUser = Depends(get_current_user),
    supabase: Client = Depends(get_supabase),
):
    """Cria um template customizado para o projeto."""
    if not Features.JURIDICO:
        raise HTTPException(status_code=404, detail="Feature não disponível")

    variables = [
        {"key": v, "label": v.replace("_", " ").title(), "type": "text"}
        for v in set(re.findall(r'\{\{(\w+)\}\}', body.get("content", "")))
    ]

    result = supabase.table("legal_templates").insert({
        "project_id": project_id,
        "category_id": body.get("category_id"),
        "name": body["name"],
        "description": body.get("description", ""),
        "content": body["content"],
        "variables": variables,
        "is_default": False,
        "created_by": user.id,
    }).execute()

    return result.data[0]


# ── IA helper ─────────────────────────────────────────────────────────────────

async def _fill_with_ai(content: str, missing_vars: list, template_name: str) -> str:
    """Usa GPT para sugerir preenchimento de campos jurídicos técnicos vazios."""
    import os
    from openai import AsyncOpenAI

    client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))

    prompt = (
        f"Você é um assistente jurídico brasileiro especializado em {template_name}.\n\n"
        f"O documento abaixo tem campos não preenchidos marcados com {{{{variavel}}}}.\n"
        f"Campos faltando: {', '.join(missing_vars)}\n\n"
        f"Preencha APENAS os campos técnicos padrão (como fundamentos jurídicos, "
        f"cláusulas padrão) com texto juridicamente correto.\n"
        f"NÃO invente dados pessoais — mantenha {{{{variavel}}}} para dados do cliente.\n\n"
        f"Documento:\n{content[:3000]}\n\n"
        f"Retorne o documento com os campos técnicos preenchidos."
    )

    response = await client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=2000,
    )

    return response.choices[0].message.content


# ── Exportação DOCX/PDF (Prompt B) ────────────────────────────────────────────

@router.get("/api/projects/{project_id}/legal/documents/{document_id}/export")
async def export_document(
    project_id: str,
    document_id: str,
    format: str = Query("docx", regex="^(docx|pdf)$"),
    user: AuthUser = Depends(get_current_user),
    supabase: Client = Depends(get_supabase),
):
    """Exporta documento jurídico como .docx ou PDF para download."""
    if not Features.JURIDICO:
        raise HTTPException(status_code=404, detail="Feature não disponível")

    result = supabase.table("legal_documents") \
        .select("*").eq("id", document_id).eq("project_id", project_id).execute()

    if not result.data:
        raise HTTPException(status_code=404, detail="Documento não encontrado")

    doc = result.data[0]
    content: str = doc["content"]
    safe_name = re.sub(r'[\\/:*?"<>|]', "-", doc["name"])

    if format == "docx":
        buf = _generate_docx(content, doc["name"])
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'},
        )
    else:
        buf = _generate_pdf(content, doc["name"])
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.pdf"'},
        )


def _generate_docx(content: str, title: str) -> io.BytesIO:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx não instalado. Execute: pip install python-docx")

    document = Document()
    section = document.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for line in content.split("\n"):
        if not line.strip():
            document.add_paragraph()
            continue
        para = document.add_paragraph(line)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = Pt(24)
        if line.isupper() or line.startswith(("EXCELENTÍSSIMO", "REQUERENTE", "REQUERIDO")):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf


def _generate_pdf(content: str, title: str) -> io.BytesIO:
    try:
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError:
        raise HTTPException(status_code=500, detail="reportlab não instalado. Execute: pip install reportlab")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=3*cm, rightMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )

    styles = getSampleStyleSheet()
    normal = ParagraphStyle("Legal", parent=styles["Normal"],
                            fontName="Times-Roman", fontSize=12, leading=24, alignment=TA_JUSTIFY)
    center = ParagraphStyle("LegalCenter", parent=normal, alignment=TA_CENTER)

    story = []
    for line in content.split("\n"):
        if not line.strip():
            story.append(Spacer(1, 12))
            continue
        s = center if line.isupper() else normal
        story.append(Paragraph(line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), s))

    doc.build(story)
    buf.seek(0)
    return buf


# ── Relatório de Atividade (Prompt D) ─────────────────────────────────────────

@router.get("/api/projects/{project_id}/legal/report")
async def get_legal_report(
    project_id: str,
    period: str = Query("30d", regex="^(7d|30d|90d)$"),
    user: AuthUser = Depends(get_current_user),
    supabase: Client = Depends(get_supabase),
):
    """Retorna relatório de atividade do projeto jurídico no período especificado."""
    if not Features.JURIDICO:
        raise HTTPException(status_code=404, detail="Feature não disponível")

    days = {"7d": 7, "30d": 30, "90d": 90}[period]
    start = (datetime.utcnow() - timedelta(days=days)).isoformat()

    docs = supabase.table("legal_documents") \
        .select("id, created_at, legal_templates(name)") \
        .eq("project_id", project_id).gte("created_at", start).execute()

    deadlines = supabase.table("legal_deadlines") \
        .select("id, status, deadline_date") \
        .eq("project_id", project_id).gte("created_at", start).execute()

    clients = supabase.table("legal_clients") \
        .select("id", count="exact").eq("project_id", project_id).execute()

    template_usage = Counter(
        (d.get("legal_templates") or {}).get("name", "—")
        for d in docs.data
    )

    daily: dict = defaultdict(int)
    for d in docs.data:
        day = d["created_at"][:10]
        daily[day] += 1

    dl_data = deadlines.data or []
    return {
        "period": period,
        "summary": {
            "documents_generated": len(docs.data),
            "deadlines_met":     len([d for d in dl_data if d["status"] == "done"]),
            "deadlines_pending": len([d for d in dl_data if d["status"] == "pending"]),
            "deadlines_overdue": len([d for d in dl_data if d["status"] == "overdue"]),
            "active_clients":    clients.count or 0,
        },
        "top_templates": [
            {"name": name, "count": count}
            for name, count in template_usage.most_common(5)
        ],
        "daily_documents": [
            {"date": k, "count": v}
            for k, v in sorted(daily.items())
        ],
    }


# ── Integração PJe/e-SAJ (Prompt E) ──────────────────────────────────────────

@router.get("/api/legal/tribunais")
async def list_tribunais(user: AuthUser = Depends(get_current_user)):
    """Lista os tribunais suportados para consulta automática."""
    from app.services.pje_scraper import TRIBUNAIS
    return {"tribunais": [
        {"key": k, "name": v["name"], "type": v["type"]}
        for k, v in TRIBUNAIS.items()
    ]}


@router.post("/api/projects/{project_id}/legal/cases/{case_id}/sync-pje")
async def sync_case_from_pje(
    project_id: str,
    case_id: str,
    body: dict,
    user: AuthUser = Depends(get_current_user),
    supabase: Client = Depends(get_supabase),
):
    """Importa prazos do PJe/e-SAJ para um processo específico."""
    if not Features.JURIDICO:
        raise HTTPException(status_code=404, detail="Feature não disponível")

    from app.services.pje_scraper import buscar_prazos_processo
    from datetime import date

    case_result = supabase.table("legal_cases").select("*") \
        .eq("id", case_id).eq("project_id", project_id).execute()

    if not case_result.data:
        raise HTTPException(status_code=404, detail="Processo não encontrado")

    case_data = case_result.data[0]
    processo = case_data.get("process_number")
    tribunal = body.get("tribunal") or case_data.get("tribunal")

    if not processo:
        raise HTTPException(status_code=400, detail="Processo sem número cadastrado")
    if not tribunal:
        raise HTTPException(status_code=400, detail="Informe o tribunal")

    result = await buscar_prazos_processo(processo, tribunal)

    if result.get("error") and not result.get("prazos"):
        raise HTTPException(status_code=400, detail=result["error"])

    imported = 0
    for prazo in result.get("prazos", []):
        existing = supabase.table("legal_deadlines").select("id") \
            .eq("project_id", project_id) \
            .eq("title", prazo["title"]) \
            .eq("deadline_date", prazo["date"]).execute()

        if not existing.data:
            supabase.table("legal_deadlines").insert({
                "project_id":    project_id,
                "case_id":       case_id,
                "title":         prazo["title"],
                "deadline_date": prazo["date"],
                "category":      prazo["category"],
                "created_by":    user.id,
            }).execute()
            imported += 1

    supabase.table("legal_cases").update({
        "last_sync": datetime.utcnow().isoformat(),
        "tribunal":  tribunal,
    }).eq("id", case_id).execute()

    return {
        "imported": imported,
        "found":    len(result.get("prazos", [])),
        "tribunal": result.get("tribunal"),
        "processo": processo,
        "warning":  result.get("error"),
    }
